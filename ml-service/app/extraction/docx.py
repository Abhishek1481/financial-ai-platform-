from __future__ import annotations

import io
import re

import docx as python_docx
from docx.table import Table as DocxTable

from app.extraction.base import ExtractionError
from app.extraction.models import ExtractedCell, ExtractedTable, ExtractionResult
from app.extraction.util import is_numeric_cell


class DocxExtractor:
    def extract(self, raw_bytes: bytes) -> ExtractionResult:
        try:
            document = python_docx.Document(io.BytesIO(raw_bytes))
        except Exception as exc:  # python-docx raises a mix of exception
            # types (PackageNotFoundError, KeyError, ...) for anything
            # that isn't a valid .docx zip package.
            raise ExtractionError(f"could not parse .docx file: {exc}") from exc

        # Body text and tables are extracted separately rather than
        # interleaved in reading order: python-docx has no simple
        # document-order iterator over both, and the structured table data
        # (headers/rows) is more useful to downstream consumers than the
        # same numbers flattened into prose.
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        text = re.sub(r"\n{3,}", "\n\n", text).strip()

        tables = [self._extract_table(t) for t in document.tables]

        return ExtractionResult(raw_text=text, tables=tables, page_count=1)

    def _extract_table(self, table: DocxTable) -> ExtractedTable:
        grid = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if not grid:
            return ExtractedTable(headers=[], rows=[])

        # DOCX has no semantic header-row markup (unlike HTML's <th>) —
        # treating the first row as a header is a convention, not
        # something the format guarantees.
        headers, *data_rows = grid
        rows = [
            [ExtractedCell(value=v, is_numeric=is_numeric_cell(v)) for v in row]
            for row in data_rows
        ]
        return ExtractedTable(headers=headers, rows=rows)
