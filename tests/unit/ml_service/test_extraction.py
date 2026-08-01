"""Unit tests for app/extraction/*. Run with the ml-service venv active:
    cd ml-service && .venv/Scripts/python.exe -m pytest ../tests/unit/ml_service -v
"""

from __future__ import annotations

import io
import sys
from pathlib import Path

import docx
import pytest

sys.path.insert(0, str(Path(__file__).parent))
from pdf_fixture import make_pdf_bytes  # noqa: E402

from app import _bootstrap  # noqa: E402,F401  (must run before generated-stub imports)
from app.extraction.base import ExtractionError  # noqa: E402
from app.extraction.docx import DocxExtractor  # noqa: E402
from app.extraction.factory import get_extractor  # noqa: E402
from app.extraction.html import HtmlExtractor  # noqa: E402
from app.extraction.pdf import PdfExtractor  # noqa: E402
from app.extraction.sec_metadata import infer_sec_metadata  # noqa: E402
from app.extraction.txt import TxtExtractor  # noqa: E402
from common.v1 import common_pb2  # noqa: E402


def make_docx_bytes(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    document = docx.Document()
    for p in paragraphs:
        document.add_paragraph(p)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for r, row_values in enumerate(table_rows):
            for c, value in enumerate(row_values):
                table.cell(r, c).text = value
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


class TestTxtExtractor:
    def test_extracts_utf8_text(self):
        result = TxtExtractor().extract("Tesla Q1 revenue: $21.3B".encode("utf-8"))
        assert result.raw_text == "Tesla Q1 revenue: $21.3B"
        assert result.page_count == 1
        assert result.tables == []

    def test_falls_back_to_latin1_when_not_valid_utf8(self):
        # 0xE9 is "é" in latin-1 but not a valid standalone UTF-8 byte, so
        # this exercises the fallback branch specifically, not just an
        # ASCII string that happens to decode fine both ways.
        raw = "café".encode("latin-1")
        with pytest.raises(UnicodeDecodeError):
            raw.decode("utf-8")  # confirms this fixture actually exercises the fallback

        result = TxtExtractor().extract(raw)
        assert result.raw_text == "café"


class TestHtmlExtractor:
    def test_extracts_text_and_strips_tags(self):
        html = b"<html><body><p>Apple's <b>Q1</b> revenue grew.</p></body></html>"
        result = HtmlExtractor().extract(html)
        assert "Apple's" in result.raw_text
        assert "Q1" in result.raw_text
        assert "<b>" not in result.raw_text

    def test_strips_script_and_style(self):
        html = b"<html><head><style>.x{color:red}</style></head><body>" \
               b"<script>alert('x')</script><p>Real content</p></body></html>"
        result = HtmlExtractor().extract(html)
        assert "alert" not in result.raw_text
        assert "color:red" not in result.raw_text
        assert "Real content" in result.raw_text

    def test_extracts_table_with_header(self):
        html = b"""
        <table>
          <caption>Revenue by segment</caption>
          <tr><th>Segment</th><th>Revenue</th></tr>
          <tr><td>Automotive</td><td>$19,963</td></tr>
        </table>
        """
        result = HtmlExtractor().extract(html)
        assert len(result.tables) == 1
        table = result.tables[0]
        assert table.headers == ["Segment", "Revenue"]
        assert table.caption == "Revenue by segment"
        assert len(table.rows) == 1
        assert table.rows[0][0].value == "Automotive"
        assert table.rows[0][0].is_numeric is False
        assert table.rows[0][1].value == "$19,963"
        assert table.rows[0][1].is_numeric is True

    def test_table_text_not_duplicated_in_raw_text(self):
        html = b"<p>Intro</p><table><tr><td>TableCellValue</td></tr></table>"
        result = HtmlExtractor().extract(html)
        assert "TableCellValue" not in result.raw_text
        assert result.tables[0].rows[0][0].value == "TableCellValue"


class TestDocxExtractor:
    def test_extracts_paragraphs(self):
        data = make_docx_bytes(["First paragraph.", "Second paragraph."])
        result = DocxExtractor().extract(data)
        assert "First paragraph." in result.raw_text
        assert "Second paragraph." in result.raw_text

    def test_extracts_table_first_row_as_header(self):
        data = make_docx_bytes(
            ["Some intro text"],
            table_rows=[["Metric", "Value"], ["Revenue", "21300"]],
        )
        result = DocxExtractor().extract(data)
        assert len(result.tables) == 1
        table = result.tables[0]
        assert table.headers == ["Metric", "Value"]
        assert table.rows[0][0].value == "Revenue"
        assert table.rows[0][1].is_numeric is True

    def test_invalid_docx_raises(self):
        with pytest.raises(ExtractionError):
            DocxExtractor().extract(b"this is not a real docx file")


class TestPdfExtractor:
    def test_extracts_text(self):
        data = make_pdf_bytes("Revenue grew 42 percent year over year")
        result = PdfExtractor().extract(data)
        assert "Revenue grew 42 percent" in result.raw_text
        assert result.page_count == 1
        assert result.tables == []  # deliberately unsupported, see pdf.py

    def test_invalid_pdf_raises(self):
        with pytest.raises(ExtractionError):
            PdfExtractor().extract(b"not a pdf at all")


class TestFactory:
    @pytest.mark.parametrize(
        "doc_type",
        [
            common_pb2.DOCUMENT_TYPE_PDF,
            common_pb2.DOCUMENT_TYPE_DOCX,
            common_pb2.DOCUMENT_TYPE_HTML,
            common_pb2.DOCUMENT_TYPE_TXT,
            common_pb2.DOCUMENT_TYPE_SEC_FILING,
        ],
    )
    def test_every_declared_doc_type_has_an_extractor(self, doc_type):
        assert get_extractor(doc_type) is not None

    def test_unknown_doc_type_raises(self):
        with pytest.raises(ValueError):
            get_extractor(common_pb2.DOCUMENT_TYPE_UNSPECIFIED)


class TestSecMetadata:
    def test_detects_10k(self):
        text = "UNITED STATES SECURITIES AND EXCHANGE COMMISSION\nFORM 10-K\nAnnual Report"
        metadata = infer_sec_metadata(text)
        assert metadata.filing_type == "10-K"

    def test_detects_10q_case_insensitive(self):
        metadata = infer_sec_metadata("this is a form 10-q quarterly filing")
        assert metadata.filing_type == "10-Q"

    def test_no_match_returns_empty(self):
        metadata = infer_sec_metadata("just some ordinary press release text")
        assert metadata.filing_type == ""

    def test_only_searches_leading_window(self):
        padding = "x" * 5000
        text = padding + " FORM 10-K "
        metadata = infer_sec_metadata(text)
        assert metadata.filing_type == ""
